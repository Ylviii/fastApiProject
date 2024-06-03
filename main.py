
import io
import os


from selenium import webdriver
#


import pandas as pd
import plotly.express as px
from fastapi.responses import JSONResponse
from starlette.responses import FileResponse
from starlette.staticfiles import StaticFiles
from fastapi import FastAPI, File, UploadFile, Form
import aspose.words as aw
import plotly.io as pio
from openpyxl import Workbook
from openpyxl.drawing.image import Image
# для pdf
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Image
from reportlab.lib import colors
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics

app = FastAPI()

app.mount("/static", StaticFiles(directory="static"), name="static")


def process_file(file_content, file_type):
    columnnames = ['Дата', 'Цена нефти, $']
    if "csv" in file_type:
        df = pd.read_csv(io.BytesIO(file_content), delimiter=';|:', engine='python')
        # Добавьте здесь обработку для CSV-файлов
    elif "excel" in file_type:
        df = pd.read_excel(io.BytesIO(file_content), names=columnnames)
        # Добавьте здесь обработку для Excel-файлов
    else:
        # Добавьте обработку для других типов файлов или выведите ошибку, если тип файла не поддерживается
        return {"error": "Unsupported file type"}

    df.dropna(inplace=True)
    month_dict = {
        'январь': 1, 'февраль': 2, 'март': 3, 'апрель': 4, 'май': 5, 'июнь': 6,
        'июль': 7, 'август': 8, 'сентябрь': 9, 'октябрь': 10, 'ноябрь': 11, 'декабрь': 12
    }

    # Преобразование названия месяца в числовой формат
    df['Дата'] = df['Дата'].str.split(', ').apply(lambda x: f"{x[0]}-{month_dict.get(x[1], x[1])}")

    # Преобразование столбца 'Месяц' в тип datetime
    df['Дата'] = pd.to_datetime(df['Дата'], format='%Y-%m')

    df['Дата'] = df['Дата'].dt.strftime('%Y-%m')
    if (file_type.__contains__('csv')):
        if (df.iloc[0]['Цена нефти, $'].__contains__(',')):
            df['Цена нефти, $'] = df['Цена нефти, $'].str.replace(',', '.').astype(float)
        else:
            df['Цена нефти, $'] = df['Цена нефти, $'].astype(float)
    return df


def AnalizRyada(file_content, file_type, period):


    k = 0.6  # Коэффициент сглаживания
    b = 0.7  # Коэффициент тренда
    p = period  # Период
    df = process_file(file_content, file_type)

    forecast_df = pd.DataFrame(columns=df.columns)
    end_date = df['Дата'].max()
    for l in range(1, p + 1):
        listT = [0]  # Значение тренда
        listY = list()  # Создаем лист значений датафрейма  с ценами на нефть
        listL = list()  # Создаем лист значений эскпоненциально сглаженного ряда
        DateList = list()  # Создаем лист дат
        kolichestvo = len(df) - 1  # Количество записей
        ### Заполнение листов с ценами на нефть и датами
        for j in range(kolichestvo, -1, -1):
            listY.append(df.iloc[j]['Цена нефти, $'])
            DateList.append(df.iloc[j]['Дата'])
        ###
        listL.append(listY[0])  # Первое значение сглаженного ряда = Первому значению ряда
        # Сохранение последней записи из датафрейма
        last_row = df.iloc[[-1]].copy()

        ### Функция добавления новой даты
        def GetDate():
            dateSplit = DateList[len(DateList) - 1].split('-')
            month = int(dateSplit[1])
            year = int(dateSplit[0])
            if (month < 12 and int(dateSplit[len(dateSplit) - 1]) > 0):
                month = month + 1
                if (month < 10):
                    monthS = '0' + str(month)
                else:
                    monthS = str(month)
            elif (month == 12):
                year = year + 1
                monthS = '01'
            return str(year) + "-" + monthS

        ###

        ### Заполнение экспоненциально-сглаженного ряда и Тренда
        for i in range(1, kolichestvo - 1):
            L = k * listY[i] + (1 - k) * (listL[i - 1] - listT[i - 1])
            listL.append(round(L, 2))
            T = b * (listL[i] - listL[i - 1]) + (1 - b) * listT[i - 1]
            listT.append(round(T, 2))
        ###

        # Добавление в ряд нового спрогнозированного значения
        new_row = pd.DataFrame(
            {'Дата': [GetDate()], 'Цена нефти, $': [round(listL[len(listL) - 1] + k * listT[len(listT) - 1], 2)]})

        # Добавление новой строки в датафрейм с помощью метода concat()
        df = pd.concat([df, new_row], ignore_index=True)

        last_row = df.iloc[-1].copy()  # Получение последней строки из исходного датафрейма df
        forecast_df = forecast_df._append(last_row, ignore_index=True)  # Добавление последней строки в forecast_df

        ### Сортировка для добавления нового значения в начало
        df['Дата'] = pd.to_datetime(df['Дата'], format='%Y-%m')
        df['Дата'] = df['Дата'].dt.strftime('%Y-%m')
        df['Цена нефти, $'] = df['Цена нефти, $'].astype(float)
        df = df.sort_values('Дата', ascending=False)
        df.reset_index(drop=True, inplace=True)


        ## Рассчитывание точности прогноза
        def TochbostAnaliza():
            dif_prognoz_and_error = 0
            for ij in range(1, len(listT)):
                prognoz_na_period = listL[ij - 1] + listT[ij - 1]
                model_error = listY[ij] - prognoz_na_period
                dif_prognoz_and_error += round((model_error * model_error) / (listY[ij] * listY[ij]), 3)

            return (1 - (dif_prognoz_and_error / (len(listY) - 1)))
        ###
    print(str(round(TochbostAnaliza(), 2) * 100) + '%')

    # print(end_date)
    # price_for_date = df[df['Дата'] == end_date]['Цена нефти, $'].astype(float)
    # price_for_date: float
    #
    # #
    # #
    # print(price_for_date)
    # asv: float = forecast_df['Цена нефти, $'].mean()
    # if price_for_date.any() > asv:
    #     print("М")
    # else:
    #     print("С")
    price_for_date = df[df['Дата'] == end_date]['Цена нефти, $'].astype(float)
    comparison_results = []
    for price in forecast_df['Цена нефти, $']:
        if price > price_for_date.values[0]:
            comparison_results.append(1)
        elif price < price_for_date.values[0]:
            comparison_results.append(-1)
        else:
            comparison_results.append(0)

    # Считаем сумму результатов
    sum_of_comparisons = sum(comparison_results)

    if sum_of_comparisons < 0:
        print("График снижается")
    elif sum_of_comparisons > 0:
        print("График повышается")
    else:
        print("График не меняется")
    # Выводим сумму
    print("Сумма результатов:", sum_of_comparisons)


    return df, end_date, sum_of_comparisons


@app.get("/")
async def root():
        return FileResponse("../fastApiProject/static/index.html")


# для принятия файла
@app.post("/file/upload-file")
async def upload_file_and_open_in_pandas(file: UploadFile = File(...)):
    file_content = await file.read()  # Чтение содержимого файла в память
    file_type = file.content_type

    df = process_file(file_content, file_type)
    fig = px.line(df, x='Дата', y='Цена нефти, $', markers=True, title='График цены нефти по датам',
                      labels={'Цена нефти, $': 'Цена нефти в долларах'})
    fig.update_layout(title_font=dict(size=20, family="Arial"), width=1285, height=550,  margin=dict(l=1, r=1, t=50, b=1),
    legend=dict(
        orientation="h",  # горизонтальная ориентация легенды
        y=1.02,  # расположение легенды немного ниже основного графика
        x=0.8,  # центрируем легенду по ширине графика
        yanchor="bottom",
        xanchor="center"
    ))  # Пример обновления шрифта заголовка графика
    # Отображение интерактивного графика без открытия в браузере

    fig_json = fig.to_json()

    return JSONResponse(content=fig_json)


@app.post("/file/analiz-file")
async def analyze_file(numericInput: int = Form(...), file: UploadFile = File(...)):
    file_content = await file.read()
    file_type = file.content_type
    period = numericInput
    ending = choose_period_ending(period)
    title = f'График с прогнозом на {period} {ending}'
    ds, end_date, sum_of_comparisons = AnalizRyada(file_content, file_type, period)

    ds[' '] = ['Исходное' if date <= end_date else 'Спрогнозированое' for date in ds['Дата']]

    figi = px.line(ds, x='Дата', y='Цена нефти, $', markers=True, title=title,
              labels={'Цена нефти, $': 'Цена нефти в долларах'}, color=' ',
              color_discrete_map={'Исходное': '#636efa', 'Спрогнозированое': 'forestgreen'}, hover_name=None)

    figi.update_layout(
        width=1285, height=550,
        margin=dict(l=1, r=1, t=1, b=1),
        legend=dict(
            orientation="h",  # горизонтальная ориентация легенды
            y=1.02,  # расположение легенды немного ниже основного графика
            x=0.5,  # центрируем легенду по ширине графика
            yanchor="bottom",
            xanchor="center"
        )
    )



    docs(title, period, ending, ds, figi,sum_of_comparisons)

    figi_json = figi.to_json()
    result = {
        "figi_json": figi.to_json(),
        "sum_of_comparisons": sum_of_comparisons
    }

    # figi.write_html("files/boxplot.html")

    return JSONResponse(content=result)


# @app.get("/file/download")
# def download_file():
#     return FileResponse(path='Отчёт.xlsx', filename='Отчёт с прогнозом.xlsx')

@app.get("/file/download/{file_format}")
def download_file(file_format: str):
    if file_format == "excel":
        file_path = "Отчёт.xlsx"
        file_name = "Отчёт с прогнозом.xlsx"
    elif file_format == "word":
        file_path = "out.docx"
        file_name = "Отчёт с прогнозом.docx"
    elif file_format == "pdf":
        file_path = "report.pdf"
        file_name = "Отчёт.pdf"
    return FileResponse(path=file_path, filename=file_name)


@app.post("/file/uploadss")
async def upload_file(file: UploadFile = File(...)):
    with open(f"{file.filename}", "wb") as f:
        f.write(await file.read())
    return {"filename": file.filename}

def choose_period_ending(count):
    if count % 10 == 1 and count % 100 != 11:
        return "период"
    elif 2 <= count % 10 <= 4 and (count % 100 < 10 or count % 100 >= 20):
        return "периода"
    else:
        return "периодов"

def docs(title, period, ending, ds, figi, sum_of_comparisons):
        figi.write_image("static/images/fig1.png")

        def sum_of_comparision_def():
            if sum_of_comparisons < 0:
                return "График снижается"
            elif sum_of_comparisons > 0:
                return "График повышается"
            else:
                return "График не меняется"





        ##############################################################################
        # Создаем новый документ

        def Create_Word():
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            document = Document()
            paragraph1 = document.add_heading()
            # Добавляем новый абзац
            paragraph1.add_run("Отчет по прогнозированию рынка нефти на " + str(period) + "\n  периода")
            # Устанавливаем шрифт, размер и цвет текста
            paragraph1.style.font.name = 'Times New Roman'
            paragraph1.style.font.size = Pt(18)
            paragraph1.style.font.color.rgb = RGBColor(0, 0, 0)
            paragraph1.style.font.bold=True
            # Устанавливаем Устанавливаем выравнивание текста
            paragraph1.alignment = 2
            paragraph1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


            # Создание таблицы
            table = document.add_table(rows=len(ds), cols=len(ds.columns))
            # Стилизация таблицы
            table.style = 'Table Grid'  # Выбор стиля таблицы

            # Изменение размеров ячеек
            for cell in table.rows[0].cells:
                cell.width = Inches(1)  # Ширина ячейки 1 дюйм

            # Изменение толщины границ
            for row in table.rows:
                for cell in row.cells:
                    for paragraph3 in cell.paragraphs:
                        for run in paragraph3.runs:
                            if run.font.color.rgb == (0, 0, 0):
                                run.font.color.rgb = (0, 0, 0)  # Цвет границ черный
                            else:
                                run.font.color.rgb = (0, 0, 0)  # Цвет текста черный
                            run.font.bold = True  # Жирный шрифт
                            run.font.size = Pt(14)  # Размер шрифта 14 пунктов

            # Заполнение таблицы данными
            for i, row in enumerate(ds.iterrows()):
                for j, col in enumerate(row[1]):
                    table.rows[i].cells[j].text = str(col)



            # Создаем объект InlinePicture из пути к файлу изображения
            image_path = 'static/images/fig1.png'
            document.add_picture(image_path, width =Inches(6))

            paragraph4 = document.add_paragraph()
            # Устанавливаем шрифт, размер и цвет текста
            paragraph4.style.font.name = 'Times New Roman'
            paragraph4.style.font.size = Pt(14)
            paragraph4.style.font.color.rgb = RGBColor(0, 0, 0)
            paragraph4.style.font.bold = False
            # Устанавливаем Устанавливаем выравнивание текста
            paragraph4.alignment = 4
            paragraph4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph4.add_run('\n      По результатам нашего пронозирования, мы сделали вывод, что '+sum_of_comparision_def()+', из-за экономических, а также политических факторов')


            paragraph2 = document.add_paragraph()
            # Устанавливаем шрифт, размер и цвет текста
            paragraph2.style.font.name = 'Times New Roman'
            paragraph2.style.font.size = Pt(14)
            paragraph2.style.font.color.rgb = RGBColor(0, 0, 0)
            paragraph2.style.font.bold = False

            # Устанавливаем Устанавливаем выравнивание текста

            paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph2.add_run("\n      Отчет создан при помощи сервиса Time Oil. Спасибо, что воспользовались нашими услугами")
            document.save('out.docx')
            #####################################################################################



        def Create_Excel():
            # create excel
            from openpyxl import Workbook
            from openpyxl.drawing.image import Image
            from openpyxl.styles import Font
            data_for_excel = ds.rename(
                columns={"Дата": "Дата", "Цена нефти, $": "Цена нефти", "Спрогнозированное": "Спрогнозированное"})
            wb = Workbook()
            ws = wb.active
            font_style = Font(name='Times New Roman', size=18, bold=True)
            ws.cell(row=1, column=6, value='Отчет по прогнозированию рынка нефти на ' + str(period) + '\n').font = font_style
            for col_num, column_name in enumerate(data_for_excel.columns, 1):
                ws.cell(row=1, column=col_num, value=column_name)

            # Запись данных
            for r_idx, row in data_for_excel.iterrows():
                for c_idx, value in enumerate(row, 1):
                    ws.cell(row=r_idx + 1, column=c_idx, value=value)

            img = Image('static/images/fig1.png')

            # Установка размеров и положения изображения в ячейке
            img.width = 1100 # ширина изображения
            img.height = 600  # высота изображения
            ws.add_image(img, 'E3')  # ячейка, в которую вставляется изображение
            ws.cell(row=35, column=5, value='По результатам нашего пронозирования, мы сделали вывод, что '+sum_of_comparision_def()+', из-за экономических, а также политических факторов')
            ws.cell(row=37, column=5, value='Отчет создан при помощи сервиса Time Oil. Спасибо, что воспользовались нашими услугами')
            # Сохранение файла Excel
            wb.save("Отчёт.xlsx")
        #######################################################################################




        #########################################################################################
        def Create_PDF():
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Image, Paragraph, Spacer
            from reportlab.lib import colors
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.pdfbase import pdfmetrics
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            # Создание объектов для PDF
            pdfmetrics.registerFont(TTFont('Arial', 'Arial.ttf'))

            # Создание объектов для PDF
            doc = SimpleDocTemplate("report.pdf", pagesize=letter)
            elements = []

            data = [ds.columns.tolist()] + ds.values.tolist()

            custom_style2 = ParagraphStyle(
                'CustomArial',
                parent=getSampleStyleSheet()["Normal"],
                fontName='Arial',
                fontSize=18,
                bold=True,
                align = 'center'
            )

            # Добавление текста после изображения
            text2 = "Отчет по прогнозированию рынка нефти на " + str(period)
            paragraphPDF = Paragraph(text2, custom_style2)
            elements.append(paragraphPDF)
            elements.append(Spacer(10, 24))
            # Создание таблицы с данными датафрейма
            t = Table(data)

            # Установка стилей для таблицы
            style = TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), 'Arial'),  # Используйте шрифт Arial
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),  # Задаем серый фон для заголовка
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),  # Задаем цвет текста для заголовка
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ])

            # Применение стилей к таблице
            t.setStyle(style)

            # Добавление таблицы в элементы документа
            elements.append(t)
            elements.append(Spacer(10, 10, 10))
            # Добавление изображения

            img = Image("static/images/fig1.png", width=600, height=200)  #

            # Добавление изображения в элементы документа
            elements.append(img)

            custom_style = ParagraphStyle(
                'CustomArial',
                parent=getSampleStyleSheet()["Normal"],
                fontName='Arial',
                fontSize=12,
                align='width'
            )

            # Добавление текста после изображения
            text = "По результатам нашего пронозирования, мы сделали вывод, что "+sum_of_comparision_def()+", из-за экономических, а также политических факторов"
            paragraph = Paragraph(text, custom_style)
            elements.append(paragraph)
            text10 = "\nОтчет создан при помощи сервиса Time Oil. Спасибо, что воспользовались нашими услугами"
            paragraph25 = Paragraph(text10, custom_style)
            elements.append(paragraph25)
            # elements.append(Spacer(10, 10, 10))
            # Генерация документа
            doc.build(elements)
            ##################################################################################
        Create_Word()
        Create_Excel()
        Create_PDF()

# # file: UploadFile = File(...)
# @app.post("/file/ds")
# async def up(file: UploadFile = File(...)):
#     file_content = await file.read()  # Чтение содержимого файла в память
#     file_type = file.content_type
#     df = process_file(file_content, file_type)
#     ds_json = df.to_json(orient='records')
#     return JSONResponse(content=ds_json)

# Призыв только 10 строк

# @app.post("/file/ds")
# async def sas(file: UploadFile = File(...)):
#     file_content = await file.read()
#     file_type = file.content_type
#     df = process_file(file_content, file_type)
#     # Ограничиваем вывод таблицы до 10 строк
#     ds_html = df.head(10).to_html(index=False)
#     return JSONResponse(content=ds_html)




@app.post("/file/daas")
async def sas(file: UploadFile = File(...)):
    file_content = await file.read()
    file_type = file.content_type
    df = process_file(file_content, file_type)

    # Добавление стилей для выравнивания таблицы по центру
    table_headers = ''.join(f'<th>{col}</th>' for col in df.columns)
    table_rows = ''.join(f'<tr>{" ".join(f"<td>{val}</td>" for val in row)}</tr>' for row in df.values)

    ds_html = f"""
      <style>
          .centered-table {{
              margin: auto;
              border-collapse: collapse;
          }}
          .centered-table th, .centered-table td {{
              border: 2px solid black;
              padding: 8px;
          }}
      </style>
      <div>
          <table class="centered-table">
              <thead>
                  <tr>{table_headers}</tr>
              </thead>
              <tbody>
                  {table_rows}
              </tbody>
          </table>
      </div>
      """

    return JSONResponse(content=ds_html)



